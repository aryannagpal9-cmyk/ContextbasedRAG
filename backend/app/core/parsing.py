import os
import pdfplumber
import docx
from bs4 import BeautifulSoup
from typing import List, Dict, Any
from pathlib import Path

class ParsedItem:
    def __init__(self, type: str, text: str, page_no: int = 1, metadata: Dict[str, Any] = None):
        self.type = type # "text", "heading", "table"
        self.text = text
        self.page_no = page_no
        self.metadata = metadata or {}

class ParsedDocument:
    def __init__(self, items: List[ParsedItem]):
        self.items = items

    def iterate_items(self):
        for item in self.items:
            # Yield (item, 0)
            yield item, 0

class DocumentParser:
    def __init__(self):
        pass

    def parse(self, file_path: str) -> ParsedDocument:
        from app.core.logging_config import logger
        logger.info(f"Initiating lightweight parse for: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        items = []

        try:
            if ext == ".pdf":
                items = self._parse_pdf(file_path)
            elif ext == ".docx":
                items = self._parse_docx(file_path)
            elif ext == ".html" or ext == ".htm":
                items = self._parse_html(file_path)
            else:
                logger.warning(f"Unsupported format: {ext}")
                items = [ParsedItem("text", "Unsupported file format.", 1)]

            logger.info(f"Successfully parsed {file_path} into {len(items)} items")
            return ParsedDocument(items)
        except Exception as e:
            logger.error(f"Parse failed for {file_path}: {e}", exc_info=True)
            raise RuntimeError(f"Error parsing document: {e}")

    def _parse_pdf(self, file_path: str) -> List[ParsedItem]:
        items = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_no = i + 1
                
                # Extract Tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        # Clean table data
                        clean_table = [[str(cell or "").strip() for cell in row] for row in table]
                        import pandas as pd
                        df = pd.DataFrame(clean_table[1:], columns=clean_table[0]) if len(clean_table) > 1 else pd.DataFrame(clean_table)
                        items.append(ParsedItem("table", df.to_markdown(index=False), page_no, {"df": df}))

                # Extract Text (excluding table areas if possible, but keep it simple for MVP)
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        if not line: continue
                        
                        # Heuristic for headings: short lines, all caps, or ending with no punctuation
                        if len(line) < 60 and (line.isupper() or not any(c in line[-1] for c in ".!?,;")):
                            items.append(ParsedItem("heading", line, page_no))
                        else:
                            items.append(ParsedItem("text", line, page_no))
        return items

    def _parse_docx(self, file_path: str) -> List[ParsedItem]:
        items = []
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            
            # Use docx style to identify headings
            if para.style.name.startswith('Heading'):
                items.append(ParsedItem("heading", text, 1))
            else:
                items.append(ParsedItem("text", text, 1))
        
        # Docx Tables
        for table in doc.tables:
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            if data:
                import pandas as pd
                df = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 else pd.DataFrame(data)
                items.append(ParsedItem("table", df.to_markdown(index=False), 1, {"df": df}))
        
        return items

    def _parse_html(self, file_path: str) -> List[ParsedItem]:
        items = []
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
            
            # Simplified: get headings and paragraphs
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table']):
                if element.name.startswith('h'):
                    items.append(ParsedItem("heading", element.get_text().strip(), 1))
                elif element.name == 'p':
                    text = element.get_text().strip()
                    if text:
                        items.append(ParsedItem("text", text, 1))
                elif element.name == 'table':
                    import pandas as pd
                    try:
                        dfs = pd.read_html(str(element))
                        if dfs:
                            items.append(ParsedItem("table", dfs[0].to_markdown(index=False), 1, {"df": dfs[0]}))
                    except:
                        continue
        return items
